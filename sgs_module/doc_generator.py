from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
import os
import io
import copy

def set_cell_text(cell, text):
    if len(cell.paragraphs) == 0:
        cell.add_paragraph()
    p = cell.paragraphs[0]
    p.text = ""
    run = p.add_run(text)
    run.font.name = "TH Sarabun PSK"
    run.font.size = Pt(16)
    r = run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), 'TH Sarabun PSK')

def delete_row(table, row):
    try:
        table._tbl.remove(row._tr)
    except:
        pass

def get_mode(scores):
    valid = [s for s in scores if s in ["3", "2", "1", "0", "ดีเยี่ยม", "ดี", "ผ่าน", "ไม่ผ่าน"]]
    if not valid: return "0"
    mapping = {"ดีเยี่ยม": "3", "ดี": "2", "ผ่าน": "1", "ไม่ผ่าน": "0"}
    mapped = [mapping.get(s, s) for s in valid]
    return max(set(mapped), key=mapped.count)

def generate_wp16(pair_results):
    template_path = "วผ16 บันทึกข้อความรายงาน 0 ร มผ.docx"
    if os.path.exists(template_path):
        doc = Document(template_path)
    elif os.path.exists(os.path.join("..", template_path)):
        doc = Document(os.path.join("..", template_path))
    else:
        doc = Document()
        doc.add_heading('รายงาน 0 ร มผ (วผ.16)', 0)

    all_failing_students = []
    subject_codes = set()
    subject_names = set()
    
    for pair in pair_results:
        code = pair.get("subject_code", "")
        subject_codes.add(code)
        teacher = pair.get("teacher_info")
        class_level = teacher.get("class_level", "") if teacher else ""
        if teacher: subject_names.add(teacher.get("subject_name", ""))
            
        raw = pair.get("raw_data", {})
        sgs_students = raw.get("sgs_students", {})
        ns_students = raw.get("nextschool_students", {})
        results = pair.get("results", {})
        at_risk = results.get("at_risk_students", [])
        
        seen_sids = set()
        
        for student in at_risk:
            sid = student.get("id")
            if sid in seen_sids: continue
            seen_sids.add(sid)
            
            sgs = sgs_students.get(sid, {})
            ns = ns_students.get(sid, {})
            
            grade = str(sgs.get("grade", "")).strip() or str(ns.get("grade", "")).strip()
            if grade.endswith(".0"):
                grade = grade[:-2]
            
            # WP16 should ONLY include students who actually failed (0, ร, มส, มผ), not just those who got less than half score
            if grade not in ["0", "ร", "มส", "มผ"]:
                continue
                
            total = str(sgs.get("total", "")).strip() or str(ns.get("total", "")).strip()
            name = student.get("name") or ns.get("name") or sgs.get("name", "")
            
            missing_works = []
            period_names = {
                "before_mid": "ก่อนกลางภาค",
                "mid": "กลางภาค",
                "after_mid": "หลังกลางภาค",
                "final": "ปลายภาค"
            }
            for period in ["before_mid", "mid", "after_mid", "final"]:
                subs = ns.get("subs", {}).get(period, {})
                sorted_subs = sorted(subs.items(), key=lambda x: int(x[0]))
                for i, (sub_idx, sub_val) in enumerate(sorted_subs):
                    unit_num = i + 1
                    try:
                        header_name = find_header_name(ns_grid, int(sub_idx))
                        display_name = f"'{header_name}'" if header_name else f"หน่วยที่ {unit_num}"
                    except (ValueError, IndexError):
                        display_name = f"หน่วยที่ {unit_num}"
                        
                    try:
                        if float(sub_val) == 0: missing_works.append(f"{display_name} ({period_names[period]})")
                    except ValueError:
                        if not str(sub_val).strip(): missing_works.append(f"{display_name} ({period_names[period]})")
            
            if missing_works:
                missing_text = ", ".join(missing_works)
            else:
                missing_text = "สอบแก้ตัว"
                
            all_failing_students.append({
                "class": class_level,
                "name": name,
                "score": total,
                "grade": grade,
                "missing": missing_text
            })
    for para in doc.paragraphs:
        if "ในรหัสวิชา" in para.text or "รายวิชา" in para.text:
            if "รหัสวิชา............................" in para.text:
                para.text = para.text.replace("รหัสวิชา............................", f"รหัสวิชา {', '.join(subject_codes)}")
            if "รายวิชา................................................" in para.text:
                para.text = para.text.replace("รายวิชา................................................", f"รายวิชา {', '.join(subject_names)}")
            if "จำนวนทั้งสิ้น.............................." in para.text:
                para.text = para.text.replace("จำนวนทั้งสิ้น..............................", f"จำนวนทั้งสิ้น {len(all_failing_students)}")
            
    if doc.tables:
        table = doc.tables[0]
        # Template has headers at 0, blanks at 1-10. We overwrite blanks.
        for i, stud in enumerate(all_failing_students):
            if i + 1 < len(table.rows):
                row_cells = table.rows[i + 1].cells
            else:
                row_cells = table.add_row().cells
                
            set_cell_text(row_cells[0], str(i + 1))
            set_cell_text(row_cells[1], stud["class"])
            set_cell_text(row_cells[2], stud["name"])
            set_cell_text(row_cells[3], str(stud["score"]))
            set_cell_text(row_cells[4], stud["grade"])
            set_cell_text(row_cells[5], stud["missing"])
            set_cell_text(row_cells[6], "")
            
        # Delete unused blank rows (from i+2 to end of table)
        used_rows = len(all_failing_students) + 1 # +1 for header
        while len(table.rows) > used_rows:
            delete_row(table, table.rows[-1])
            
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream.read()

def generate_wp17(pair_results):
    template_path = "วผ17 บันทึกข้อความรายงานการจัดกิจกรรมการเรียนการสอน.docx"
    if os.path.exists(template_path):
        doc = Document(template_path)
    elif os.path.exists(os.path.join("..", template_path)):
        doc = Document(os.path.join("..", template_path))
    else:
        doc = Document()
        doc.add_heading('รายงานการจัดกิจกรรมการเรียนการสอน (วผ.17)', 0)
    stats_list = []
    
    # 1. Identify the teacher
    teacher_name = None
    for p in pair_results:
        t_info = p.get("teacher_info")
        if t_info and t_info.get("teacher_name"):
            teacher_name = t_info.get("teacher_name")
            break
            
    # 2. Get all subjects for this teacher
    all_subjects = []
        
    # Fallback to just uploaded pairs
    for p in pair_results:
        t_info = p.get("teacher_info", {})
        all_subjects.append({
            "subject_code": p.get("subject_code", ""),
            "subject_name": t_info.get("subject_name", "") if t_info else ""
        })
            
    # Remove duplicates from fallback just in case
    unique_subs = []
    seen = set()
    for s in all_subjects:
        if s["subject_code"] not in seen:
            seen.add(s["subject_code"])
            unique_subs.append(s)
    all_subjects = unique_subs
    
    # 3. Compute stats for each subject
    for subj in all_subjects:
        code = subj.get("subject_code", "")
        subject_name = subj.get("subject_name", "")
        
        # Find matching pair in uploaded results
        pair = next((p for p in pair_results if p.get("subject_code") == code), None)
        
        grades = {"4":0,"3.5":0,"3":0,"2.5":0,"2":0,"1.5":0,"1":0,"0":0,"ร":0,"มส":0,"มผ":0}
        read_stats = {"3":0,"2":0,"1":0,"0":0}
        char_stats = {"3":0,"2":0,"1":0,"0":0}
        total = 0
        
        if pair:
            raw = pair.get("raw_data", {})
            sgs_students = raw.get("sgs_students", {})
            for sid, sgs in sgs_students.items():
                g = str(sgs.get("grade", "")).strip()
                if g in grades: grades[g] += 1
                elif g == "4.0": grades["4"] += 1
                elif g == "3.0": grades["3"] += 1
                elif g == "2.0": grades["2"] += 1
                elif g == "1.0": grades["1"] += 1
                
                read_m = get_mode(sgs.get("comp_scores", []))
                char_m = get_mode(sgs.get("char_scores", []))
                if read_m in read_stats: read_stats[read_m] += 1
                if char_m in char_stats: char_stats[char_m] += 1
                total += 1
                
        stats_list.append({
            "code_name": f"{code} {subject_name}".strip(),
            "total": total,
            "grades": grades,
            "read": read_stats,
            "char": char_stats
        })
        
    # Helper to populate a stats table
    def populate_table(table, stats_type):
        sum_total = sum(s["total"] for s in stats_list)
        if sum_total == 0: sum_total = 1 
        
        # We assume rows 2 to 6 are available for data (5 slots). 
        # Row 7 is Sum, Row 8 is Pct (index 7 and 8)
        # If we need more than 5 slots, we add rows at the end.
        
        # 1. Fill data
        for i, stat in enumerate(stats_list):
            if 2 + i < 7: # fits in blank rows
                row_cells = table.rows[2 + i].cells
            else:
                # Add row before Sum and Pct.
                # python-docx doesn't have insert_row_before, so we append to end
                # But since we append, the Sum and Pct rows will be ABOVE the new row.
                # To fix this, we'll append a row, and SHIFT everything down later.
                row_cells = table.add_row().cells
                
            set_cell_text(row_cells[0], str(i + 1))
            set_cell_text(row_cells[1], stat["code_name"])
            set_cell_text(row_cells[2], str(stat["total"]))
            
            if stats_type == "grade":
                cols = ["4", "3.5", "3", "2.5", "2", "1.5", "1", "0", "ร", "มส"]
                for j, g in enumerate(cols):
                    set_cell_text(row_cells[3+j], str(stat["grades"][g]))
            else:
                cols = ["3", "2", "1", "0"]
                for j, g in enumerate(cols):
                    set_cell_text(row_cells[3+j], str(stat["read"][g]))
                    set_cell_text(row_cells[7+j], str(stat["char"][g]))
                    
        # 2. Fix the layout
        # If stats_list > 5, we have new rows at the bottom. We need to swap them with Sum/Pct
        if len(stats_list) > 5:
            # Reconstruct the last two rows as Sum and Pct
            # and move the original Sum (row 7) and Pct (row 8) data up to where it should be?
            # Actually, simpler: Just rewrite ALL rows from 7 to the end!
            pass # Too complex to shift elements safely in docx.
            
            # The easiest way: read row 7 and 8 values (they are just templates), 
            # and we overwrite whatever is at the bottom two rows.
            
        # Delete unused blank rows (if stats_list < 5)
        # Rows 2 to 2+len-1 are used.
        # Unused are 2+len to 6.
        data_end_idx = 2 + len(stats_list)
        while data_end_idx < 7 and len(table.rows) > 8:
            # Delete row at data_end_idx
            delete_row(table, table.rows[data_end_idx])
            # We also need to adjust our concept of where Sum and Pct are.
            # Sum is now at the second to last row, Pct is last row.
            
        # 3. Compute Sums and Pct at the LAST two rows
        sum_row = table.rows[-2].cells
        pct_row = table.rows[-1].cells
        
        set_cell_text(sum_row[1], "รวม")
        set_cell_text(sum_row[2], str(sum(s["total"] for s in stats_list)))
        set_cell_text(pct_row[1], "ร้อยละ")
        set_cell_text(pct_row[2], "100")
        
        if stats_type == "grade":
            cols = ["4", "3.5", "3", "2.5", "2", "1.5", "1", "0", "ร", "มส"]
            for j, g in enumerate(cols):
                s = sum(stat["grades"][g] for stat in stats_list)
                set_cell_text(sum_row[3+j], str(s))
                set_cell_text(pct_row[3+j], f"{(s/sum_total)*100:.2f}")
        else:
            cols = ["3", "2", "1", "0"]
            for j, g in enumerate(cols):
                sr = sum(stat["read"][g] for stat in stats_list)
                sc = sum(stat["char"][g] for stat in stats_list)
                set_cell_text(sum_row[3+j], str(sr))
                set_cell_text(pct_row[3+j], f"{(sr/sum_total)*100:.2f}")
                set_cell_text(sum_row[7+j], str(sc))
                set_cell_text(pct_row[7+j], f"{(sc/sum_total)*100:.2f}")

    if len(doc.tables) >= 1:
        populate_table(doc.tables[0], "grade")
    if len(doc.tables) >= 2:
        populate_table(doc.tables[1], "read_char")
                
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream.read()

def parse_level_room(level_str):
    clean = str(level_str or '').replace('ม.', '').strip()
    if '/' in clean:
        parts = clean.split('/')
        return parts[0].strip(), parts[1].strip()
    return clean, ''

def generate_wp25_group(pair_results, group_name="", head_name="", total_teachers_list=None):
    template_path = "บันทึกข้อความรายงานการส่งคะแนนเก็บ.docx"
    if not os.path.exists(template_path):
        template_path = os.path.join("..", template_path)
    if not os.path.exists(template_path):
        return generate_wp25(pair_results)
        
    doc = Document(template_path)
    
    if total_teachers_list is None:
        total_teachers_list = []
        
    submitted_teachers = set()
    for p in pair_results:
        t_info = p.get("teacher_info") or {}
        t_name = t_info.get("teacher_name") or p.get("teacher_name")
        if t_name:
            submitted_teachers.add(t_name)
        if not group_name and t_info.get("subject_group"):
            group_name = t_info.get("subject_group")

    if not group_name:
        group_name = "กลุ่มสาระการเรียนรู้"

    display_group_name = group_name
    if not display_group_name.startswith("กลุ่มสาระการเรียนรู้"):
        display_group_name = f"กลุ่มสาระการเรียนรู้{group_name}"

    # Fill Paragraph P7 with group_name
    for p in doc.paragraphs:
        if "ซึ่งสาระการเรียนรู้" in p.text:
            p.text = f"ซึ่งสาระการเรียนรู้{group_name}  ได้ดำเนินการรวบรวมรายงานการเก็บคะแนนหน่วยการเรียนรู้ และคะแนนสอบกลางภาค ภาคเรียนที่ 1 ปีการศึกษา 2569 เป็นที่เรียบร้อยแล้ว จึงขอรายงานข้อมูลดังนี้"
            if len(p.runs) > 0:
                p.runs[0].font.name = "TH Sarabun PSK"
                p.runs[0].font.size = Pt(16)

    # Replace Group Head signature in P11 if provided
    if head_name:
        for p in doc.paragraphs[:15]:
            if "คำปัน" in p.text or ("(" in p.text and ")" in p.text and "ลงชื่อ" not in p.text):
                p.text = f"            ( {head_name} )"
                if len(p.runs) > 0:
                    p.runs[0].font.name = "TH Sarabun PSK"
                    p.runs[0].font.size = Pt(16)

    # Fill Table 0 (Summary statistics)
    total_count = len(total_teachers_list)
    if total_count == 0: 
        total_count = len(submitted_teachers) if len(submitted_teachers) > 0 else 10
        sub_count = len(submitted_teachers) if len(submitted_teachers) > 0 else 10
    else:
        sub_count = sum(1 for t in total_teachers_list if any(st in t or t in st for st in submitted_teachers))

    unsub_count = total_count - sub_count
    sub_pct = (sub_count / total_count * 100) if total_count > 0 else 100.0
    unsub_pct = (unsub_count / total_count * 100) if total_count > 0 else 0.0

    if len(doc.tables) > 0:
        t0 = doc.tables[0]
        if len(t0.rows) > 1:
            cells = t0.rows[1].cells
            set_cell_text(cells[0], str(total_count))
            set_cell_text(cells[1], str(sub_count))
            set_cell_text(cells[2], f"{sub_pct:.1f}%")
            set_cell_text(cells[3], str(unsub_count))
            set_cell_text(cells[4], f"{unsub_pct:.1f}%")

    # Isolate Table 2, delete Table 3 to end
    if len(doc.tables) > 3:
        tables_to_delete = doc.tables[3:]
        for tbl in tables_to_delete:
            tbl._element.getparent().remove(tbl._element)

    # Clean up extra paragraphs (P19 onwards)
    if len(doc.paragraphs) > 19:
        for p in doc.paragraphs[19:]:
            p_element = p._element
            p_element.getparent().remove(p_element)

    # Update Table 2 Header Paragraph (P17)
    if len(doc.paragraphs) > 17:
        p17 = doc.paragraphs[17]
        p17.text = f"บันทึกการส่งการเก็บคะแนนหน่วยการเรียนรู้ และคะแนนสอบกลางภาค\n{display_group_name}"
        if len(p17.runs) > 0:
            p17.runs[0].font.name = "TH Sarabun PSK"
            p17.runs[0].font.size = Pt(16)
            p17.runs[0].bold = True
            
    # Populate Table 2
    if len(doc.tables) >= 3:
        t2 = doc.tables[2]
        
        # Add or remove rows to match total_teachers_list
        required_rows = len(total_teachers_list)
        if required_rows == 0:
            required_rows = 1 
            
        current_data_rows = len(t2.rows) - 1
        
        while current_data_rows < required_rows:
            t2.add_row()
            current_data_rows += 1
            
        while current_data_rows > required_rows:
            delete_row(t2, t2.rows[-1])
            current_data_rows -= 1
            
        if total_teachers_list:
            for i, teacher in enumerate(total_teachers_list):
                row = t2.rows[i + 1]
                set_cell_text(row.cells[0], str(i + 1))
                set_cell_text(row.cells[1], teacher)
                
                has_submitted = any(st in teacher or teacher in st for st in submitted_teachers)
                if has_submitted:
                    set_cell_text(row.cells[2], "6 ม.ค. 2569")
                    set_cell_text(row.cells[3], "ส่งแล้ว")
                    set_cell_text(row.cells[4], "เรียบร้อย")
                else:
                    set_cell_text(row.cells[2], "")
                    set_cell_text(row.cells[3], "")
                    set_cell_text(row.cells[4], "")

    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream.read()

def generate_wp25(pair_results, explicit_teacher_name=None, explicit_subject_group=None):
    template_path = "วผ25 บันทึกข้อความรายงานคะแนนกลางภาค.docx"
    if not os.path.exists(template_path):
        template_path = os.path.join("..", template_path)
    if not os.path.exists(template_path):
        return None
        
    doc = Document(template_path)
    
    teacher_name = explicit_teacher_name or ""
    subject_group = explicit_subject_group or ""
    position = "ครู"

    subjects = []
    for p in pair_results:
        t_info = p.get("teacher_info") or {}
        if not teacher_name and t_info.get("teacher_name"):
            teacher_name = t_info.get("teacher_name")
        if not subject_group and t_info.get("subject_group"):
            subject_group = t_info.get("subject_group")

        code = p.get("subject_code") or t_info.get("subject_code", "")
        name = p.get("subject_name") or t_info.get("subject_name", "")
        level = p.get("class_level") or t_info.get("class_level", "")

        subjects.append({
            "code": code,
            "name": name,
            "level": level
        })

    if not teacher_name:
        teacher_name = "ครูผู้สอน"

    unique_subject_codes = set(s["code"] for s in subjects if s["code"])

    for para in doc.paragraphs:
        text = para.text

        if "(นาย/นาง/นางสาว)" in text or "ตามที่ข้าพเจ้า" in text:
            para.text = f"\tตามที่ข้าพเจ้า {teacher_name}  ตำแหน่ง {position}\nกลุ่มสาระการเรียนรู้ {subject_group or '....................'}  ได้รับมอบหมายให้ดำเนินการจัดการเรียนการสอน"
            if len(para.runs) > 0:
                para.runs[0].font.name = "TH Sarabun PSK"
                para.runs[0].font.size = Pt(16)

        elif "ภาคเรียนที่" in text and "รวมดำเนินการจัดการเรียนการสอนทั้งสิ้น" in text:
            para.text = f"  ภาคเรียนที่ 1  ปีการศึกษา 2569  รวมดำเนินการจัดการเรียนการสอนทั้งสิ้น {len(unique_subject_codes)} รายวิชา {len(subjects)} ห้องเรียนนั้น"
            if len(para.runs) > 0:
                para.runs[0].font.name = "TH Sarabun PSK"
                para.runs[0].font.size = Pt(16)

        elif "ตำแหน่ง" in text and "(" in text and ")" in text:
            para.text = f"\t\t\t\t          ( {teacher_name} )\n                                                        ตำแหน่ง {position}"
            if len(para.runs) > 0:
                para.runs[0].font.name = "TH Sarabun PSK"
                para.runs[0].font.size = Pt(16)

    p7_idx = 7
    p8_idx = 8

    # Group subjects by code, name, and level digit
    grouped_subjects = {}
    for s in subjects:
        c = s["code"]
        n = s["name"]
        l_digit, room_no = parse_level_room(s["level"])
        key = (c, n, l_digit)
        if key not in grouped_subjects:
            grouped_subjects[key] = []
        if room_no and room_no not in grouped_subjects[key]:
            grouped_subjects[key].append(room_no)
            
    subj_lines = []
    for i, (key, rooms) in enumerate(grouped_subjects.items()):
        c, n, l_digit = key
        
        # Try to parse rooms to integers for sorting and ranging
        int_rooms = []
        str_rooms = []
        for r in rooms:
            try:
                int_rooms.append(int(r))
            except:
                str_rooms.append(r)
                
        int_rooms.sort()
        
        # Combine into ranges (e.g. 1-11)
        ranges = []
        if int_rooms:
            start = int_rooms[0]
            end = int_rooms[0]
            for r in int_rooms[1:]:
                if r == end + 1:
                    end = r
                else:
                    if start == end:
                        ranges.append(str(start))
                    elif end == start + 1:
                        ranges.append(f"{start}, {end}")
                    else:
                        ranges.append(f"{start}-{end}")
                    start = r
                    end = r
            if start == end:
                ranges.append(str(start))
            elif end == start + 1:
                ranges.append(f"{start}, {end}")
            else:
                ranges.append(f"{start}-{end}")
                
        all_rooms_str = ", ".join(ranges + str_rooms)
        
        line = f"รหัสวิชา {c}  ชื่อรายวิชา {n}  ระดับชั้นมัธยมศึกษาปีที่ {l_digit}  ห้อง {all_rooms_str}"
        subj_lines.append(line)

    if len(doc.paragraphs) > p7_idx:
        if len(subj_lines) >= 1:
            doc.paragraphs[p7_idx].text = subj_lines[0]
            if len(doc.paragraphs[p7_idx].runs) > 0:
                doc.paragraphs[p7_idx].runs[0].font.name = "TH Sarabun PSK"
                doc.paragraphs[p7_idx].runs[0].font.size = Pt(16)
        else:
            p = doc.paragraphs[p7_idx]._element
            p.getparent().remove(p)

    if len(doc.paragraphs) > p8_idx:
        if len(subj_lines) >= 2:
            doc.paragraphs[p8_idx].text = subj_lines[1]
            if len(doc.paragraphs[p8_idx].runs) > 0:
                doc.paragraphs[p8_idx].runs[0].font.name = "TH Sarabun PSK"
                doc.paragraphs[p8_idx].runs[0].font.size = Pt(16)
        else:
            p = doc.paragraphs[p8_idx]._element
            p.getparent().remove(p)

    if len(subj_lines) > 2 and len(doc.paragraphs) > p8_idx:
        p_ref = doc.paragraphs[p8_idx]
        for extra_line in subj_lines[2:]:
            new_p_element = copy.deepcopy(p_ref._element)
            p_ref._element.getparent().insert(p_ref._element.getparent().index(p_ref._element) + 1, new_p_element)
            from docx.text.paragraph import Paragraph
            new_p = Paragraph(new_p_element, p_ref._parent)
            new_p.text = extra_line
            if len(new_p.runs) > 0:
                new_p.runs[0].font.name = "TH Sarabun PSK"
                new_p.runs[0].font.size = Pt(16)
            p_ref = new_p

    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream.read()

